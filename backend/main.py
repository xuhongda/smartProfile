from fastapi import FastAPI, Request, Depends, UploadFile, File, HTTPException, staticfiles
from typing import List
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
import time
import os
from database import get_db, Document
from core.services import container
from core.models.api_response import APIResponse
from core.services.hybrid_search_service import hybrid_search_service
from core.services.vector_store import vector_store

# 从容器获取服务和配置
config = container.resolve('config')
database_service = container.resolve('database_service')
tokenizer_service = container.resolve('tokenizer_service')
file_parser_service = container.resolve('file_parser_service')
document_manager_service = container.resolve('document_manager_service')
search_service = container.resolve('search_service')
generic_ai_client = container.resolve('generic_ai_client')
model_manager_service = container.resolve('model_manager_service')

# 创建上传目录
UPLOAD_DIR = config.get('upload.directory')
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

# 初始化数据库
database_service.init_database()

app = FastAPI(
    title=config.get('api.title'),
    description=config.get('api.description'),
    version=config.get('api.version')
)

# 挂载静态文件服务
app.mount("/uploads", staticfiles.StaticFiles(directory=UPLOAD_DIR), name="uploads")

# 挂载 models 目录作为静态文件服务
MODELS_DIR = os.path.join(os.path.dirname(__file__), 'models')
if not os.path.exists(MODELS_DIR):
    os.makedirs(MODELS_DIR)
app.mount("/models", staticfiles.StaticFiles(directory=MODELS_DIR), name="models")

# 配置 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=config.get('cors.allow_origins'),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 自定义中间件 - 日志中间件
@app.middleware("http")
async def log_requests(request: Request, call_next):
    start_time = time.time()
    response = await call_next(request)
    process_time = time.time() - start_time
    response.headers["X-Process-Time"] = str(process_time)
    print(f"[{request.method}] {request.url.path} - {process_time:.4f}s")
    return response

@app.get("/")
async def root():
    return {"message": "Unstructured Data Search API"}

@app.get("/health")
async def health_check():
    """健康检查接口"""
    return {
        "status": "healthy",
        "message": "API 服务运行正常",
        "timestamp": time.time()
    }

@app.get("/api/model/path")
async def get_model_path():
    """获取模型文件夹路径
    
    Returns:
        模型文件夹路径
    """
    try:
        # 获取 models 文件夹的绝对路径
        models_path = os.path.abspath(MODELS_DIR)
        return {
            "success": True,
            "path": models_path
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"获取模型文件夹路径失败: {str(e)}"
        }

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """文件上传与解析接口
    
    Args:
        file: 上传的文件
        db: 数据库会话
    
    Returns:
        上传和解析结果
    """
    try:
        # 检查文件大小
        contents = await file.read()
        max_size = config.get('upload.max_size')
        if len(contents) > max_size:
            raise HTTPException(status_code=413, detail=f"文件大小超过限制（最大{max_size // 1024 // 1024}MB）")
        
        # 解析文件
        parse_result = file_parser_service.parse_file(contents, file.filename)
        
        # 对中文内容进行分词处理
        if tokenizer_service.is_chinese(parse_result.content):
            content = tokenizer_service.tokenize_text(parse_result.content)
        else:
            content = parse_result.content
        
        # 保存原文件到服务器文件系统
        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, "wb") as f:
            f.write(contents)
        
        # 创建或更新文档
        doc, message = document_manager_service.create_document(
            db, file.filename, parse_result.file_type, content
        )
        
        # 添加到 ChromaDB 向量数据库
        try:
            metadata = {
                "filename": doc.filename,
                "file_type": doc.file_type,
                "created_at": str(doc.created_at),
                "updated_at": str(doc.updated_at)
            }
            vector_store.add_document(doc.uuid, content, metadata)
            print(f"文档已添加到 ChromaDB: {doc.filename}")
        except Exception as e:
            print(f"添加到 ChromaDB 失败: {str(e)}")
        
        return {
            "success": True,
            "message": message,
            "file_uuid": doc.uuid,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "content_length": doc.content_length,
            "file_path": f"/uploads/{doc.filename}"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")

@app.get("/search")
async def search(
    q: str,
    page: int = 1,
    page_size: int = 10,
    enable_ai: bool = False,
    db: Session = Depends(get_db)
):
    """全文搜索接口
    
    Args:
        q: 搜索关键词
        page: 页码
        page_size: 每页大小
        enable_ai: 是否启用AI增强
        db: 数据库会话
    
    Returns:
        搜索结果
    """
    try:
        if not q:
            raise HTTPException(status_code=400, detail="搜索关键词不能为空")
        
        # 检查是否使用 @AI 标记
        use_ai = enable_ai or q.strip().startswith("@AI")
        if use_ai and q.strip().startswith("@AI"):
            # 移除 @AI 标记
            q = q.strip()[3:].strip()
        
        # 普通搜索：先使用原来的搜索服务获取所有相关结果
        search_result = search_service.search(db, q, page, page_size)
        result_dict = search_result.to_dict()
        results = result_dict.get("results", [])
        
        # 如果启用AI增强，将搜索结果发给AI生成回答
        if use_ai:
            # 只有当有搜索结果时才生成AI响应，避免基于错误数据的虚假回答
            if results:
                # 生成AI响应
                search_results_for_ai = [{"filename": r.get("filename"), "content": r.get("content")} for r in results]
                ai_response = generic_ai_client.generate_ai_response(q, search_results_for_ai)
                if ai_response:
                    # 将AI回答放在搜索结果后面
                    result_dict["ai_response"] = ai_response
            else:
                # 如果没有搜索结果，不调用AI，避免虚假回答
                print("没有搜索结果，不调用AI生成回答")
        
        return {
            "success": True,
            **result_dict
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")

@app.get("/documents")
async def get_documents(
    page: int = 1,
    page_size: int = 10,
    db: Session = Depends(get_db)
):
    """获取文档列表
    
    Args:
        page: 页码
        page_size: 每页大小
        db: 数据库会话
    
    Returns:
        文档列表
    """
    try:
        # 调用文档管理服务
        doc_list = document_manager_service.get_documents(db, page, page_size)
        
        # 转换为API响应格式
        result_dict = doc_list.to_dict()
        return {
            "success": True,
            **result_dict
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取文档列表失败: {str(e)}")

@app.delete("/documents/{doc_uuid}")
async def delete_document(
    doc_uuid: str,
    db: Session = Depends(get_db)
):
    """删除文档
    
    Args:
        doc_uuid: 文档UUID
        db: 数据库会话
    
    Returns:
        删除结果
    """
    try:
        # 查找文档
        doc = document_manager_service.get_document_by_uuid(db, doc_uuid)
        
        # 删除文件系统中的文件
        file_path = os.path.join(UPLOAD_DIR, doc.filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # 删除文档
        document_manager_service.delete_document_by_uuid(db, doc_uuid)
        
        return {
            "success": True,
            "message": f"文档 '{doc.filename}' 删除成功"
        }
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除文档失败: {str(e)}")

@app.get("/db-check")
async def db_check(db: Session = Depends(get_db)):
    """数据库连接检查接口"""
    try:
        # 测试数据库连接
        from sqlalchemy import text
        db.execute(text("SELECT 1")).scalar()
        # 检查文档表是否存在
        result = db.execute(text("SELECT name FROM sqlite_master WHERE type='table' AND name='documents'")).fetchone()
        if not result:
            return {
                "success": False,
                "message": "文档表不存在"
            }
        # 检查FTS表是否存在
        result_fts = db.execute(text("SELECT name FROM sqlite_master WHERE type='table' AND name='documents_fts'")).fetchone()
        if not result_fts:
            return {
                "success": False,
                "message": "FTS表不存在"
            }
        return {
            "success": True,
            "message": "数据库连接正常，表结构完整"
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"数据库连接失败: {str(e)}"
        }

@app.get("/preview/{doc_uuid}")
async def preview_document(
    doc_uuid: str,
    db: Session = Depends(get_db)
):
    """文件预览接口
    
    Args:
        doc_uuid: 文档UUID
        db: 数据库会话
    
    Returns:
        预览数据
    """
    try:
        # 查找文档
        doc = document_manager_service.get_document_by_uuid(db, doc_uuid)
        
        # 构建文件路径
        file_path = os.path.join(UPLOAD_DIR, doc.filename)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取文件内容
        with open(file_path, "rb") as f:
            file_content = f.read()
        
        # 根据文件类型处理预览数据
        if doc.file_type == "word":
            # 解析Word文档为HTML
            from docx import Document
            from io import BytesIO
            
            doc_obj = Document(BytesIO(file_content))
            html_content = "<html><body>"
            
            for para in doc_obj.paragraphs:
                if para.text:
                    html_content += f"<p>{para.text}</p>"
            
            # 处理表格
            for table in doc_obj.tables:
                html_content += "<table border='1' style='border-collapse: collapse;'>"
                for row in table.rows:
                    html_content += "<tr>"
                    for cell in row.cells:
                        html_content += f"<td>{cell.text}</td>"
                    html_content += "</tr>"
                html_content += "</table><br>"
            
            html_content += "</body></html>"
            
            return {
                "success": True,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "content": html_content,
                "content_type": "text/html"
            }
        
        elif doc.file_type == "excel":
            # 解析Excel文档为JSON
            import pandas as pd
            from io import BytesIO
            import json
            
            df = pd.read_excel(BytesIO(file_content), engine='openpyxl')
            # 处理nan值
            df = df.fillna('')
            # 转换为JSON
            json_data = df.to_dict(orient='records')
            
            return {
                "success": True,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "content": json_data,
                "content_type": "application/json"
            }
        
        elif doc.file_type == "txt":
            # 直接返回文本内容
            import chardet
            result = chardet.detect(file_content)
            encoding = result['encoding'] or 'utf-8'
            text_content = file_content.decode(encoding)
            
            return {
                "success": True,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "content": text_content,
                "content_type": "text/plain"
            }
        
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型: {doc.file_type}")
        
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件预览失败: {str(e)}")

@app.post("/api/ai/config")
async def set_ai_config(
    request: Request
):
    """设置AI服务配置
    
    Args:
        request: 包含配置数据的请求
    
    Returns:
        配置结果
    """
    try:
        # 解析请求体
        config_data = await request.json()
        
        # 更新配置
        if "api_base_url" in config_data:
            config.set("ai_service.api_base_url", config_data["api_base_url"])
        if "api_key" in config_data:
            config.set("ai_service.api_key", config_data["api_key"])
        if "embedding_model" in config_data:
            config.set("ai_service.embedding_model", config_data["embedding_model"])
        if "speech_model" in config_data:
            config.set("ai_service.speech_model", config_data["speech_model"])
        if "chat_model" in config_data:
            config.set("ai_service.chat_model", config_data["chat_model"])
        if "timeout" in config_data:
            config.set("ai_service.timeout", config_data["timeout"])
        
        # 更新AI客户端配置
        generic_ai_client.update_config()
        
        # 保存配置到文件
        config.save()
        
        return {
            "success": True,
            "message": "AI服务配置更新成功"
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"配置更新失败: {str(e)}"
        }

@app.post("/api/ai/test-connection")
async def test_ai_connection(
    request: Request
):
    """测试AI服务连接
    
    Args:
        request: 包含API地址和密钥的请求
    
    Returns:
        连接测试结果
    """
    try:
        # 解析请求体
        data = await request.json()
        api_base_url = data.get("api_base_url", config.get("ai_service.api_base_url"))
        api_key = data.get("api_key", config.get("ai_service.api_key"))
        
        if not api_base_url:
            return {
                "success": False,
                "message": "API地址不能为空"
            }
        
        # 测试连接
        import httpx
        client = httpx.Client(timeout=httpx.Timeout(5.0))
        headers = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        
        # 尝试调用模型列表接口
        response = client.get(f"{api_base_url.rstrip('/')}/models", headers=headers)
        
        if response.status_code == 200:
            return {
                "success": True,
                "message": "连接测试成功"
            }
        else:
            return {
                "success": False,
                "message": f"连接测试失败，状态码: {response.status_code}"
            }
    except httpx.ConnectError:
        return {
            "success": False,
            "message": "连接失败，请检查API地址是否正确"
        }
    except httpx.TimeoutException:
        return {
            "success": False,
            "message": "连接超时，请检查服务是否正常运行"
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"连接测试失败: {str(e)}"
        }

@app.post("/api/model/upload")
async def upload_model(
    files: List[UploadFile] = File(...)
):
    """上传模型
    
    Args:
        files: 上传的模型文件列表
    
    Returns:
        上传结果
    """
    try:
        # 调用模型管理服务上传模型
        result = await model_manager_service.upload_model(files)
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        return {
            "success": False,
            "message": f"上传模型失败: {str(e)}"
        }

@app.get("/api/model/list")
async def list_models():
    """获取模型列表
    
    Returns:
        模型列表
    """
    try:
        models = model_manager_service.get_models()
        return {
            "success": True,
            "models": models
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"获取模型列表失败: {str(e)}"
        }

@app.delete("/api/model/{model_name}")
async def delete_model(
    model_name: str
):
    """删除模型
    
    Args:
        model_name: 模型名称
    
    Returns:
        删除结果
    """
    try:
        # 调用模型管理服务删除模型
        result = model_manager_service.delete_model(model_name)
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        return {
            "success": False,
            "message": f"删除模型失败: {str(e)}"
        }



@app.post("/api/embedding")
async def create_embedding(
    request: Request
):
    """生成文本嵌入
    
    Args:
        request: 包含文本数据的请求
    
    Returns:
        嵌入向量
    """
    try:
        # 解析请求体
        data = await request.json()
        text = data.get("text", "")
        
        if not text:
            return {
                "success": False,
                "message": "文本不能为空"
            }
        
        # 生成嵌入
        embedding = generic_ai_client.create_embedding(text)
        
        if embedding:
            return {
                "success": True,
                "embedding": embedding
            }
        else:
            return {
                "success": False,
                "message": "生成嵌入失败，请检查模型配置"
            }
    except Exception as e:
        return {
            "success": False,
            "message": f"生成嵌入失败: {str(e)}"
        }

@app.post("/api/speech/recognize")
async def speech_recognize(
    file: UploadFile = File(...)
):
    """语音识别
    
    Args:
        file: 上传的音频文件
    
    Returns:
        识别结果
    """
    try:
        # 读取音频文件
        audio_data = await file.read()
        
        if not audio_data:
            return {
                "success": False,
                "message": "音频文件不能为空"
            }
        
        # 进行语音识别
        result = generic_ai_client.speech_recognition(audio_data)
        
        if result:
            return {
                "success": True,
                "text": result
            }
        else:
            return {
                "success": False,
                "message": "语音识别失败，请检查模型配置"
            }
    except Exception as e:
        return {
            "success": False,
            "message": f"语音识别失败: {str(e)}"
        }
